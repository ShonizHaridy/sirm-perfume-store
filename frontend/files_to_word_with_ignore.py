import os
from docx import Document

def write_files_to_word(directory, doc, ignored_directories, ignored_files):
    # Walk through the directory
    for root, dirs, files in os.walk(directory):
        # Remove ignored directories from the list of directories to traverse
        dirs[:] = [d for d in dirs if os.path.join(root, d) not in ignored_directories]
        
        for file in files:
            file_path = os.path.join(root, file)
            
            # Skip ignored files
            if file_path in ignored_files or file in ignored_files:
                continue
                
            # Add file path to the Word document
            doc.add_paragraph(f"File Path: {file_path}")

            # Read the file content
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    doc.add_paragraph("Content:")
                    doc.add_paragraph(content)
            except Exception as e:
                doc.add_paragraph(f"Could not read file {file_path}: {e}")

            # Add a page break after each file's content
            doc.add_page_break()

def get_ignored_items(directory, item_type):
    items = []
    count = input(f"Enter number of {item_type} to ignore (0 for none): ")
    try:
        count = int(count)
        for i in range(count):
            name = input(f"Enter {item_type} name {i+1}: ")
            if item_type == "directories":
                full_path = os.path.join(directory, name)
                items.append(full_path)
            else:  # files
                # Add both possible paths (full path and just filename)
                full_path = os.path.join(directory, name)
                items.extend([full_path, name])
    except ValueError:
        print(f"Invalid number, ignoring {item_type}.")
    return items

def main():
    # Specify the directory you want to scan
    directory = input("Enter the directory path: ")

    # Get directories and files to ignore
    ignored_directories = get_ignored_items(directory, "directories")
    ignored_files = get_ignored_items(directory, "files")

    # Create a new Word document
    doc = Document()
    doc.add_heading('File Paths and Contents', level=1)

    # Write the files and contents to the document
    write_files_to_word(directory, doc, ignored_directories, ignored_files)

    # Save the document
    output_file = os.path.join(directory, 'file_contents.docx')
    doc.save(output_file)
    print(f"Document saved as: {output_file}")

if __name__ == "__main__":
    main()
